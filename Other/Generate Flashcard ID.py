import docx
from docx.enum.section import WD_SECTION
from docx.shared import Pt
import string
import random
from io import BytesIO
import os
import time



test_doc = "C:\\Users\\Stephan\\OneDrive\\Learning System\\flashcards template.docx"
dest_dir = "C:\\Users\\Stephan\\OneDrive\\Flashcards\\Statistical Experimental Design\\Montgomery - Design and Analysis of Experiments - 6th\\Flashcards Word Files"


def generateCode():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))


if __name__ == "__main__":
    # section_list = ["01.01", "01.02", "01.03", "01.04", "01.05", "01.06", "01.07", "01.08", "01.09", "01.10",
    #                 "02.01", "02.02", "02.03", "02.04", "02.05", "02.06", "02.07", "02.08",
    #                 "03.01", "03.02", "03.03", "03.04", "03.05", "03.06",
    #                 "04.01", "04.02", "04.03", "04.04", "04.05", "04.06", "04.07", "04.08", "04.09", "04.10",
    #                 "05.01", "05.02", "05.03", "05.04",
    #                 "06.01", "06.02", "06.03", "06.04", "06.05", "06.06",
    #                 "07.01", "07.02", "07.03", "07.04", "07.05", "07.06", "07.07", "07.08", "07.09",
    #                 "08.03", "08.04", "08.05",
    #                 "09.01", "09.02", "09.03", "09.04", "09.05", "09.06", "09.07", "09.08", "09.09",
    #                 "10.01", "10.02", "10.03", "10.04", "10.05", "10.06", "10.07", "10.08", "10.09",
    #                 "11.01", "11.02", "11.03", "11.04",
    #                 ]
    #for section in section_list:
    section = "02.00"
    if os.path.exists(os.path.join(dest_dir, section, section + ".docx")):
        print(section + ".docx already exists.")
    else:
        doc = docx.Document(test_doc)
        doc.styles['Normal'].font.size = Pt(8)
        n = 200
        flashcard_number = 1
        for i in range(n):
            #id = generateCode()
            doc.sections[-1].header.paragraphs[0].text = "Front"
            #doc.sections[-1].footer.paragraphs[0].text = str("\t"*2 + id)
            doc.add_page_break()
            doc.add_section(WD_SECTION.NEW_PAGE)
            doc.sections[-1].footer.is_linked_to_previous = False
            doc.sections[-1].header.is_linked_to_previous = False
            doc.sections[-1].header.paragraphs[0].text = "Back"
            doc.sections[-1].footer.paragraphs[0].text = str("\t"*4 + str(flashcard_number))
            if i < n - 1:
                doc.add_page_break()
                doc.add_section(WD_SECTION.NEW_PAGE)
                doc.sections[-1].footer.is_linked_to_previous = False
                doc.sections[-1].header.is_linked_to_previous = False
            flashcard_number += 1
        doc.save(os.path.join(dest_dir, section, section + ".docx"))

    # for i in range(5):
    #     id = generateCode()
    #     doc.sections[-1].footer.is_linked_to_previous = False1
    #     doc.sections[-1].header.is_linked_to_previous = False
    #     doc.sections[-1].header.paragraphs[0].text = "Answer"
    #     doc.sections[-1].footer.paragraphs[0].text = str("\t"*2 + id)
    #     doc.add_section(WD_SECTION.NEW_PAGE)
    #     doc.add_page_break()
    #     doc.sections[-1].footer.is_linked_to_previous = False
    #     doc.sections[-1].header.is_linked_to_previous = False
    #     doc.sections[-1].header.paragraphs[0].text = "Question"
    #     doc.sections[-1].footer.paragraphs[0].text = str("\t"*2 + id)
    #     if i < 4:
    #         doc.add_section(WD_SECTION.NEW_PAGE)
    #         doc.add_page_break()
    # doc.save(test_out_doc)
    # os.startfile(test_out_doc)
