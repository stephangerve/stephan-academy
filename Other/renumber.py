import docx


doc_path = "C:\\Users\\Stephan\\OneDrive\\Flashcards\\Statistical Experimental Design\\Montgomery - Design and Analysis of Experiments - 6th\\Flashcards Word Files\\04.00\\04.00.docx"


if __name__ == "__main__":
    doc = docx.Document(doc_path)
    flashcard_number = 1
    for page in doc.sections:
        if page.footer.paragraphs[0].text.split('\t\t')[-1] != '':
            page.footer.paragraphs[0].text = str("\t"*4 + str(flashcard_number))
            flashcard_number += 1
    doc.save(doc_path)